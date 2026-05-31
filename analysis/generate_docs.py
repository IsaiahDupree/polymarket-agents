"""
Generate two Word documents from the PolymarketAutomation codebase + live DB:

  1. analysis/agent-variables-and-strategies.docx
       Reference manual: every genome kind, every parameter (with bounds),
       every env var that affects training/promotion, the fitness formula,
       the promotion gates, and a cross-reference matrix.

  2. analysis/architecture-and-timing.docx
       Operating report: system overview, component breakdown, event
       cadences, recent factory output, win-rate distribution, top agents,
       strategic positioning of strategies + gap analysis. Includes
       matplotlib charts (PNG, embedded).

Reads from:
  - src/lib/arena/genome.ts (param bounds — hardcoded mirror below since
    parsing TS in Python is fragile; update if genome.ts changes)
  - data/polymarket.db (live stats: alive agents, recent campaigns)
  - data/factory-state.json (current PID/uptime per factory)

Run:
  python analysis/generate_docs.py

Both files land in analysis/. Charts in analysis/charts/.
"""
from __future__ import annotations

import json
import sqlite3
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")  # no GUI backend — headless
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths

REPO = Path(__file__).resolve().parent.parent
ANALYSIS = REPO / "analysis"
CHARTS = ANALYSIS / "charts"
DB = REPO / "data" / "polymarket.db"
STATE = REPO / "data" / "factory-state.json"
OUT_VARS = ANALYSIS / "agent-variables-and-strategies.docx"
OUT_ARCH = ANALYSIS / "architecture-and-timing.docx"
OUT_LATENCY = ANALYSIS / "latency-and-event-map.docx"

ANALYSIS.mkdir(exist_ok=True)
CHARTS.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# Strategy metadata (mirrors src/lib/arena/genome.ts and arena/sim.ts)

STRATEGIES: list[dict[str, Any]] = [
    {
        "kind": "poly_short_binary_directional",
        "title": "Polymarket Short Binary Directional",
        "venue": "Polymarket",
        "summary": (
            "Directional trades on Polymarket rolling 5/15-min binary contracts (BTC Up/Down). "
            "Uses Coinbase price velocity over a short window to predict whether YES will resolve "
            "above the entry threshold. The 'consistent-winner' BTC-5m factory targets this kind."
        ),
        "params": {
            "vel_window_min":          ([1, 5],          "Coinbase velocity measurement window in minutes."),
            "vel_entry_pct":           ([0.0001, 0.02],  "Min |price change| as fraction over the window to fire an entry."),
            "pre_cutoff_min":          ([2, 4],          "Reject entries closer than N minutes before resolution."),
            "max_window_min":          ([3, 16],         "Max age in minutes of the binary contract being traded."),
            "max_yes_price_for_buy":   ([0.40, 0.85],    "Don't BUY YES above this price (avoid expensive entries)."),
            "min_yes_price_for_sell":  ([0.15, 0.60],    "Don't SELL YES below this price (avoid selling cheap)."),
            "entry_size_usd":          ([1, 50],         "Stake per entry. Default $2 under staged-stake phase 1."),
            "max_positions_per_asset": ([1, 4],          "Cap concurrent positions on the same crypto asset."),
        },
        "fires_when": "Velocity threshold crossed AND inside window AND YES price in band.",
        "live_eligible": True,
    },
    {
        "kind": "poly_fade_spike",
        "title": "Polymarket Fade Spike",
        "venue": "Polymarket",
        "summary": (
            "Fade short-term overreactions. When YES jumps `threshold_pts` over `lookback_h` followed by "
            "`confirm_quiet_h` of low activity, take the opposite side expecting mean reversion."
        ),
        "params": {
            "threshold_pts":     ([3, 10],   "Required magnitude of spike in price-cents to qualify."),
            "lookback_h":        ([6, 72],   "Window over which to measure the spike."),
            "confirm_quiet_h":   ([2, 24],   "Quiet period after the spike before fading."),
            "entry_size_usd":    ([5, 100],  "Stake per entry."),
            "exit_target_pts":   ([1, 8],    "Take-profit in price-cents."),
            "stop_pts":          ([2, 10],   "Stop-loss in price-cents."),
            "time_stop_h":       ([12, 168], "Exit unconditionally after N hours."),
        },
        "fires_when": "spike >= threshold AND quiet period passed since spike.",
        "live_eligible": False,
        "note": "Visible-book dependent — fills unreliably on FAK live orders; sim-only.",
    },
    {
        "kind": "poly_breakout",
        "title": "Polymarket Breakout",
        "venue": "Polymarket",
        "summary": (
            "Ride breakouts above the recent N-period high. Long when YES breaks above the prior range "
            "by a multiplier; exit on target / stop / time-stop."
        ),
        "params": {
            "lookback_h":      ([6, 168],   "Range-detection window."),
            "breakout_mult":   ([1.05, 3.0], "Multiplier above the lookback high required to fire."),
            "entry_size_usd":  ([5, 100],   "Stake per entry."),
            "target_pts":      ([1, 15],    "Profit target."),
            "stop_pts":        ([2, 15],    "Stop-loss."),
            "time_stop_h":     ([12, 168],  "Time-based exit."),
        },
        "fires_when": "YES price exceeds (lookback_high * breakout_mult).",
        "live_eligible": False,
        "note": "Visible-book dependent.",
    },
    {
        "kind": "cb_breakout",
        "title": "Coinbase Breakout",
        "venue": "Coinbase",
        "summary": (
            "Range breakout on Coinbase spot (BTC/ETH/SOL). Long when price breaks above the rolling "
            "lookback high by a configurable multiplier; exits on percentage target / stop / time."
        ),
        "params": {
            "product_id":      (["BTC-USD", "ETH-USD", "SOL-USD"], "Coinbase market."),
            "lookback_min":    ([15, 1440],     "Range-detection window in minutes."),
            "breakout_mult":   ([1.01, 1.10],   "Above-high multiplier (1.05 = 5% break)."),
            "entry_size_usd":  ([5, 100],       "Stake per entry."),
            "target_pct":      ([0.001, 0.10],  "Profit target as fraction (0.05 = +5%)."),
            "stop_pct":        ([0.002, 0.10],  "Stop-loss as fraction."),
            "time_stop_min":   ([30, 4320],     "Time-based exit."),
        },
        "fires_when": "price > rolling_high * breakout_mult.",
        "live_eligible": True,
    },
    {
        "kind": "cb_mean_reversion",
        "title": "Coinbase Mean Reversion",
        "venue": "Coinbase",
        "summary": (
            "Buy oversold dips on Coinbase. Enters when z-score below `z_entry` (1-2.5 sigma below mean), "
            "exits when z reverts to `z_exit` or hits stop/time."
        ),
        "params": {
            "product_id":     (["BTC-USD", "ETH-USD", "SOL-USD"], "Coinbase market."),
            "lookback_min":   ([60, 4320],    "Rolling window for mean/sigma."),
            "z_entry":        ([1.0, 2.5],   "Std-dev threshold to enter (positive = N sigma below)."),
            "z_exit":         ([-1.0, 1.0],  "Std-dev threshold to exit."),
            "entry_size_usd": ([5, 100],     "Stake per entry."),
            "stop_pct":       ([0.005, 0.10], "Pct-based stop loss."),
            "time_stop_min":  ([30, 4320],   "Time-based exit."),
        },
        "fires_when": "z-score < -z_entry (sufficiently oversold).",
        "live_eligible": True,
    },
    {
        "kind": "cb_momentum_burst",
        "title": "Coinbase Momentum Burst",
        "venue": "Coinbase",
        "summary": (
            "Trade velocity/acceleration spikes. Enter when price change exceeds `vel_entry_pct` over "
            "`vel_window_min` with sufficient acceleration. Short or long depending on direction bias."
        ),
        "params": {
            "product_id":      (["BTC-USD", "ETH-USD", "SOL-USD"], "Coinbase market."),
            "vel_window_min":  ([3, 20],       "Velocity window in minutes."),
            "vel_entry_pct":   ([0.001, 0.012], "Min pct change over the window."),
            "accel_min":       ([0.00005, 0.003], "Min acceleration (2nd derivative) to fire."),
            "entry_size_usd":  ([5, 100],      "Stake per entry."),
            "target_pct":      ([0.001, 0.02], "Profit target."),
            "stop_pct":        ([0.002, 0.02], "Stop-loss."),
            "time_stop_min":   ([5, 120],      "Time-based exit."),
            "direction_bias":  (["long_only", "long_short"], "Long-only or both directions."),
        },
        "fires_when": "|velocity| > vel_entry_pct AND |accel| > accel_min.",
        "live_eligible": True,
    },
    {
        "kind": "cross_venue_arb",
        "title": "Cross-Venue Arbitrage",
        "venue": "Polymarket + Coinbase",
        "summary": (
            "Fade the spread between Polymarket-implied probability and Coinbase-implied probability "
            "via Black-Scholes vol. Enters when |edge_pts| > threshold, exits when spread normalizes."
        ),
        "params": {
            "cb_product_id":      (["BTC-USD", "ETH-USD"], "Coinbase pair."),
            "poly_condition_id":  ([], "Specific Polymarket market — provided by seed list."),
            "edge_pts":           ([2, 20],   "Required edge in price-cents to enter."),
            "bs_vol_window_days": ([7, 60],   "Vol estimation window for the BS-implied prob."),
            "entry_size_usd":     ([5, 100],  "Stake per entry."),
            "time_stop_h":        ([12, 168], "Time-based exit."),
        },
        "fires_when": "|poly_price - bs_implied_prob|*100 > edge_pts.",
        "live_eligible": False,
        "note": "Visible-book dependent; sim-only.",
    },
    {
        "kind": "polymarket_market_maker",
        "title": "Polymarket Market Maker",
        "venue": "Polymarket",
        "summary": (
            "Post bid AND ask around the midpoint, collect the spread. No directional view. "
            "Sim-only for now; the executor doesn't post true two-sided quotes yet."
        ),
        "params": {
            "spread_pts":     ([0.5, 5],   "Half-spread in price-cents from midpoint."),
            "stop_pts":       ([1, 10],    "Inventory stop-loss."),
            "time_stop_h":    ([1, 12],    "Cap inventory holding period."),
            "entry_size_usd": ([1, 20],    "Quote size."),
        },
        "fires_when": "Continuously — quotes both sides while inside risk caps.",
        "live_eligible": True,
        "note": "Live-eligible kind list includes it but the executor is sim-only.",
    },
    {
        "kind": "llm_probability_oracle",
        "title": "LLM Probability Oracle",
        "venue": "Polymarket",
        "summary": (
            "Calls an LLM (Claude family) to estimate the true probability of each market's YES resolving. "
            "Applies EV + Kelly rails: only enters when EV >= min_ev_pct."
        ),
        "params": {
            "model":              (["claude-opus-4-7", "claude-sonnet-4-6", "claude-haiku-4-5-20251001"], "Claude model."),
            "min_ev_pct":         ([0.05, 0.20],  "Minimum positive EV to fire (5-20%)."),
            "max_calls_per_tick": ([1, 5],        "Rate-limit LLM calls per arena tick."),
            "cache_ttl_min":      ([5, 240],      "Memoize probabilities for N minutes."),
            "entry_size_usd":     ([5, 100],      "Stake per entry."),
        },
        "fires_when": "P(true) - market_price > min_ev_pct.",
        "live_eligible": True,
    },
    {
        "kind": "random_walk_baseline",
        "title": "Random Walk Baseline",
        "venue": "Polymarket",
        "summary": (
            "Null hypothesis control — fires random buys/sells. If a real strategy can't outperform "
            "this baseline on the same fitness function, the strategy has no edge."
        ),
        "params": {
            "trade_prob":     ([0.001, 0.10], "Probability of trading per tick."),
            "buy_bias_pct":   ([0.30, 0.70],  "Fraction of trades that go BUY."),
            "entry_size_usd": ([5, 50],       "Stake per entry."),
        },
        "fires_when": "Random number < trade_prob.",
        "live_eligible": False,
    },
    {
        "kind": "category_specialist",
        "title": "Category Specialist",
        "venue": "Polymarket",
        "summary": (
            "Applies fade_spike or breakout but only to markets tagged in a single category "
            "(geopolitics, crypto, sports, etc). Tests whether category-restricted strategies "
            "outperform the asset-class-wide version."
        ),
        "params": {
            "category":        (["geopolitics", "elections", "crypto", "sports", "macro", "weather", "tech", "other"], "Category filter."),
            "inner_strategy":  (["fade_spike", "breakout"], "Which strategy to apply inside the category."),
            "threshold_pts":   ([3, 12],   "Forwarded to inner strategy."),
            "lookback_h":      ([6, 72],   "Forwarded to inner strategy."),
            "confirm_quiet_h": ([2, 24],   "Forwarded to inner strategy."),
            "entry_size_usd":  ([5, 100],  "Stake per entry."),
            "exit_target_pts": ([1, 8],    "Target."),
            "stop_pts":        ([2, 10],   "Stop."),
            "time_stop_h":     ([12, 168], "Time exit."),
            "breakout_mult":   ([1.05, 2.5], "Forwarded if inner_strategy=breakout."),
        },
        "fires_when": "Inner strategy fires AND market.category == category.",
        "live_eligible": False,
    },
    {
        "kind": "wallet_copy_filtered",
        "title": "Wallet Copy (Filtered)",
        "venue": "Polymarket",
        "summary": (
            "Mirror a tracked wallet's trades, but only within a category, scaled by size_pct_of_source. "
            "Has min source win-rate / trade-count gates so the copy only triggers on proven sources."
        ),
        "params": {
            "copy_category":       (["geopolitics", "elections", "crypto", "sports", "macro", "weather", "tech", "other"], "Only copy trades in this category."),
            "size_pct_of_source":  ([0.001, 0.10], "Fraction of source's notional."),
            "max_size_usd":        ([1, 100],     "Hard cap on each copy."),
            "delay_min":           ([1, 60],      "Wait N minutes after source's trade (avoid race)."),
            "min_source_win_rate": ([0.40, 0.90], "Don't copy a wallet below this win rate."),
            "min_source_trades":   ([5, 200],     "Don't copy until source has N closed trades."),
        },
        "fires_when": "Tracked wallet trades + filters pass + delay elapsed.",
        "live_eligible": False,
    },
    {
        "kind": "multi_strategy",
        "title": "Multi-Strategy Composite",
        "venue": "Multiple",
        "summary": (
            "Composite that wraps 2-4 sub-strategies. Selection mode 'priority' picks the FIRST sub "
            "that returns a non-HOLD decision. Evolution mutates the sub list + each sub's params."
        ),
        "params": {
            "selection":      (["priority"],  "Picker: first non-HOLD wins."),
            "entry_size_usd": ([5, 100],      "Forwarded to whichever sub fires."),
        },
        "fires_when": "Any sub fires (in priority order).",
        "live_eligible": "conditional",
        "note": "Live-eligible only if at least one sub is in the live-eligible kind list.",
    },
    {
        "kind": "poly_arbitrage_set",
        "title": "Polymarket Arbitrage Set",
        "venue": "Polymarket",
        "summary": (
            "Buy YES + NO together when their asks (plus fees) sum below $1. Risk-free locked profit "
            "at resolution. Ported from polymarket-2dollar-bot/polybot/microstructure.py."
        ),
        "params": {
            "min_edge":       ([0.001, 0.05],  "Required (1 - cost - fees)/cost edge."),
            "max_set_cost":   ([0.85, 0.999],  "Cap on YES_ask + NO_ask."),
            "fee_bps":        ([0, 200],       "Polymarket taker fee in bps. Default 100."),
            "entry_size_usd": ([2, 100],       "Per-leg stake."),
        },
        "fires_when": "YES_ask + NO_ask + fees < $1 by min_edge.",
        "live_eligible": True,
        "note": "Rare opportunity; capacity is low.",
    },
    {
        "kind": "poly_repricing",
        "title": "Polymarket Repricing",
        "venue": "Polymarket + Coinbase",
        "summary": (
            "Directional bet on the lag between Coinbase spot and Polymarket's reprice. Computes fair "
            "P(YES) from spot vs strike via tanh-scaled BS rule, fires when |fair - market| >= min_edge. "
            "Uses the same 4 event-timing gates as markov_persistence."
        ),
        "params": {
            "min_edge":                    ([0.02, 0.20],    "Required |fair - market| edge."),
            "max_yes_price_for_buy":       ([0.40, 0.85],    "Don't BUY YES above this."),
            "min_yes_price_for_sell":      ([0.15, 0.60],    "Don't SELL YES below this."),
            "entry_size_usd":              ([2, 100],        "Per-entry stake."),
            "min_time_to_resolution_min":  ([0, 30],         "Skip if too close to expiry."),
            "max_time_to_resolution_min":  ([1, 999],        "Skip if too far out."),
            "event_phase_filter":          (["any", "opening", "mid-window", "late-window", "mid-or-late", "tradeable"], "Lifecycle phase gate."),
            "max_signal_age_sec":          ([1, 9999],       "Max age of latest Coinbase WS tick."),
        },
        "fires_when": "Spot moved + Polymarket hasn't repriced + signal is fresh + phase OK.",
        "live_eligible": True,
        "note": "Most timing-sensitive of the new kinds.",
    },
    {
        "kind": "poly_directional_arb_tilt",
        "title": "Polymarket Directional Arb Tilt",
        "venue": "Polymarket",
        "summary": (
            "Arb base (YES + NO < $1) AND a velocity-based model view → tilt toward the under-priced "
            "side. Asymmetric exposure but the arb floor bounds downside."
        ),
        "params": {
            "min_edge":         ([0.001, 0.05], "Required arb-set edge as a floor."),
            "max_set_cost":     ([0.85, 0.999], "Cap on YES + NO ask sum."),
            "fee_bps":          ([0, 200],      "Polymarket taker fee."),
            "model_window_min": ([1, 30],       "Window for velocity-based directional model."),
            "entry_size_usd":   ([2, 100],      "Per-leg stake."),
        },
        "fires_when": "Arb base AND velocity gives a directional view.",
        "live_eligible": True,
    },
    {
        "kind": "poly_near_resolution",
        "title": "Polymarket Near-Resolution Scrape",
        "venue": "Polymarket",
        "summary": (
            "Buy a near-certain side trading at 0.95-0.99 in the final minutes. High win rate / small "
            "reward (the '$2 → ~$0.30' profile). Time-to-resolution gated to bound tail risk."
        ),
        "params": {
            "min_price":         ([0.85, 0.99],   "Min YES price to qualify."),
            "max_price":         ([0.95, 0.999],  "Max YES price (above = thin upside)."),
            "max_seconds_left":  ([15, 600],      "Max seconds-to-resolution for the entry."),
            "entry_size_usd":    ([2, 50],        "Per-entry stake."),
        },
        "fires_when": "YES price in [min, max] AND seconds-to-expiry <= max_seconds_left.",
        "live_eligible": True,
    },
]


# ---------------------------------------------------------------------------
# Env var reference (these affect training/promotion/factory behavior)

ENV_VARS: list[dict[str, str]] = [
    # Risk budget
    {"name": "RISK_STAKE_USD",                  "default": "2",      "purpose": "Per-trade dollar stake (single USD anchor)."},
    {"name": "RISK_N_AGENTS",                   "default": "3",      "purpose": "Number of live capsules to share the budget."},
    {"name": "RISK_DAILY_STAKES_AT_RISK",       "default": "1",      "purpose": "Losing stakes per day before a capsule pauses for the day."},
    {"name": "RISK_LIFETIME_STAKES_AT_RISK",    "default": "2",      "purpose": "Lifetime losing stakes before a capsule permanently pauses."},
    {"name": "RISK_FILL_RATE_HEADROOM",         "default": "200",    "purpose": "Multiplier on per-capsule trade-count cap (200 = 200 trades/day allowed)."},
    {"name": "MIN_LIVE_CAPSULE_PNL_USD",        "default": "96",     "purpose": "Lifetime PnL gate on the live-promotion boundary."},

    # Fitness mode
    {"name": "ARENA_FITNESS_MODE",              "default": "winrate", "purpose": "Either 'pnl_dd' (legacy) or 'winrate' (default, biased toward 90%+)."},
    {"name": "ARENA_WINRATE_POWER",             "default": "2",      "purpose": "Exponent on win_rate term — 2 steepens gradient near 1.0."},
    {"name": "ARENA_MIN_TRADES_FOR_RANKING",    "default": "30",     "purpose": "Min closed trades for an agent to qualify for ranking (else sentinel)."},

    # Graduation gates (sim → paper)
    {"name": "GRADUATION_MIN_PNL_USD",          "default": "10",     "purpose": "Min lifetime PnL to graduate sim→paper."},
    {"name": "GRADUATION_MIN_TRADES",           "default": "15",     "purpose": "Min closed trades to graduate."},
    {"name": "GRADUATION_MIN_WIN_RATE",         "default": "0.90",   "purpose": "Min win rate to graduate (matches the BTC 5m Up/Down goal)."},
    {"name": "GRADUATION_AUTO_STAGE_CAPITAL",   "default": "50",     "purpose": "Starting capital for a graduated paper capsule."},
    {"name": "GRADUATION_AUTO_STAGE_MAX_DAILY_LOSS",  "default": "25", "purpose": "Daily-loss cap on the graduated capsule."},
    {"name": "GRADUATION_AUTO_STAGE_MAX_TOTAL_DRAWDOWN", "default": "10", "purpose": "Total-drawdown cap on the graduated capsule."},

    # Auto-promote (paper → live)
    {"name": "ALLOW_AUTO_PROMOTE",              "default": "(unset)", "purpose": "Required '1' to enable any auto-promotion to live."},
    {"name": "ARENA_AUTO_PROMOTE_MIN_TRADES",   "default": "3",      "purpose": "Min trades to qualify for live promotion."},
    {"name": "ARENA_AUTO_PROMOTE_MIN_WIN_RATE", "default": "0.90",   "purpose": "Min win rate on the live boundary. Set 0 to disable."},
    {"name": "ARENA_AUTO_PROMOTE_LIVE_KINDS",   "default": "(see kinds list)", "purpose": "CSV of live-eligible kinds (safety ceiling)."},
    {"name": "ARENA_AUTO_PROMOTE_MIN_SHARE",    "default": "0.15",   "purpose": "Min capital share floor per elite agent."},
    {"name": "DYNAMIC_KIND_BLACKLIST",          "default": "1",      "purpose": "If '1', rolling-window perf query auto-blacklists negative-PnL kinds."},
    {"name": "ARENA_DYNAMIC_KIND_WINDOW_DAYS",  "default": "30",     "purpose": "Window size for dynamic blacklist evaluation."},

    # Factory cadences
    {"name": "BARS_PER_STEP",                   "default": "(N/A on TS side)", "purpose": "Step granularity in arena tick (TradingBot Python only)."},
    {"name": "FACTORY_MULTI_KINDS",             "default": "(all 12 non-BTC kinds)", "purpose": "CSV of genome kinds the multi-factory should cover."},
    {"name": "FACTORY_MULTI_FAST_VARIANTS",     "default": "8",      "purpose": "Per-kind variant count for fast cycles."},
    {"name": "FACTORY_MULTI_DEEP_VARIANTS",     "default": "15",     "purpose": "Per-kind variant count for deep cycles."},
    {"name": "FACTORY_MULTI_FAST_SEED",         "default": "2",      "purpose": "Per-kind top-K to seed from fast cycles."},
    {"name": "FACTORY_MULTI_DEEP_SEED",         "default": "3",      "purpose": "Per-kind top-K to seed from deep cycles."},
    {"name": "FACTORY_DRY_RUN",                 "default": "(unset)", "purpose": "If '1', factory plans cycles but doesn't write to DB."},
    {"name": "FACTORY_DASHBOARD_REFRESH_MS",    "default": "5000",   "purpose": "Dashboard refresh interval in milliseconds."},

    # Campaign scoring
    {"name": "CAMPAIGN_SCORE_FN",               "default": "composite", "purpose": "Campaign ranking function: 'composite' (PnL+trades+winrate-DD) or alternatives."},
    {"name": "CAMPAIGN_W_PNL",                  "default": "(see code)", "purpose": "Weight on PnL in composite score."},
    {"name": "CAMPAIGN_W_TRADES",               "default": "(see code)", "purpose": "Weight on trade count."},
    {"name": "CAMPAIGN_W_WIN_RATE",             "default": "(see code)", "purpose": "Weight on win rate."},
    {"name": "CAMPAIGN_W_DD",                   "default": "(see code)", "purpose": "Penalty weight on drawdown."},
    {"name": "CAMPAIGN_FORCE_ENTRY_SIZE_USD",   "default": "2 (factory:btc-5m)", "purpose": "Override every variant's entry_size_usd to this."},
    {"name": "CAMPAIGN_FORCE_PRICE_BAND",       "default": "1 (factory:btc-5m)", "purpose": "Force narrower YES price band for consistent-winner profile."},
]


# ---------------------------------------------------------------------------
# DB queries

def fetch_live_stats() -> dict[str, Any]:
    """Pull live snapshot from polymarket.db. Returns empty dict if DB absent."""
    if not DB.exists():
        return {"db_missing": True}
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    try:
        totals = conn.execute(
            "SELECT alive, COUNT(*) AS n, SUM(trades_count) AS tt, SUM(realized_pnl_usd) AS pp "
            "FROM paper_agents GROUP BY alive"
        ).fetchall()
        alive_count = next((r["n"] for r in totals if r["alive"] == 1), 0)
        dead_count = next((r["n"] for r in totals if r["alive"] == 0), 0)

        top = conn.execute(
            "SELECT id, name, COALESCE(json_extract(genome_json, '$.kind'), 'unknown') AS kind, "
            "trades_count, wins_count, realized_pnl_usd "
            "FROM paper_agents WHERE alive=1 AND trades_count >= 30 "
            "ORDER BY (CAST(wins_count AS REAL)/trades_count) DESC, realized_pnl_usd DESC "
            "LIMIT 10"
        ).fetchall()

        # win-rate distribution among alive >= 30 trades
        dist = {"90-100%": 0, "80-90%": 0, "70-80%": 0, "60-70%": 0, "50-60%": 0, "<50%": 0}
        qual = conn.execute(
            "SELECT wins_count, trades_count FROM paper_agents "
            "WHERE alive=1 AND trades_count >= 30"
        ).fetchall()
        for r in qual:
            wr = r["wins_count"] / r["trades_count"]
            if wr >= 0.90: dist["90-100%"] += 1
            elif wr >= 0.80: dist["80-90%"] += 1
            elif wr >= 0.70: dist["70-80%"] += 1
            elif wr >= 0.60: dist["60-70%"] += 1
            elif wr >= 0.50: dist["50-60%"] += 1
            else: dist["<50%"] += 1

        # per-kind alive count
        per_kind = conn.execute(
            "SELECT COALESCE(json_extract(genome_json, '$.kind'), 'unknown') AS kind, "
            "COUNT(*) AS n, AVG(trades_count) AS avg_t, AVG(realized_pnl_usd) AS avg_pnl "
            "FROM paper_agents WHERE alive=1 GROUP BY kind ORDER BY n DESC"
        ).fetchall()

        # recent campaigns
        camps = conn.execute(
            "SELECT name, kind, candidates_produced, best_pnl_usd, created_at "
            "FROM training_campaigns WHERE created_at >= datetime('now','-24 hours') "
            "ORDER BY id DESC LIMIT 30"
        ).fetchall()

        # snapshot coverage
        snap_total = conn.execute("SELECT COUNT(*) AS n FROM market_snapshots").fetchone()["n"]
        snap_window = conn.execute(
            "SELECT MIN(captured_at) AS lo, MAX(captured_at) AS hi FROM market_snapshots"
        ).fetchone()

        return {
            "alive_count": alive_count,
            "dead_count": dead_count,
            "qualifying_count": len(qual),
            "best_win_rate": (max((r["wins_count"] / r["trades_count"] for r in qual), default=0.0)),
            "win_rate_distribution": dist,
            "top_agents": [dict(r) for r in top],
            "per_kind": [dict(r) for r in per_kind],
            "recent_campaigns": [dict(r) for r in camps],
            "snapshot_total": snap_total,
            "snapshot_earliest": snap_window["lo"],
            "snapshot_latest": snap_window["hi"],
        }
    finally:
        conn.close()


def read_factory_state() -> dict[str, Any]:
    if not STATE.exists():
        return {"missing": True}
    return json.loads(STATE.read_text())


# ---------------------------------------------------------------------------
# Charts (matplotlib → PNG → embedded in docx)

def chart_win_rate_distribution(dist: dict[str, int], path: Path) -> None:
    """Bar chart of win-rate buckets, top-down (90-100 first)."""
    labels = ["90-100%", "80-90%", "70-80%", "60-70%", "50-60%", "<50%"]
    counts = [dist.get(k, 0) for k in labels]
    colors = ["#2ecc71", "#27ae60", "#f1c40f", "#e67e22", "#e74c3c", "#7f8c8d"]
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.barh(labels, counts, color=colors)
    ax.set_xlabel("Alive agents (≥30 closed trades)")
    ax.set_title("Win-rate distribution across qualifying alive agents")
    ax.invert_yaxis()  # 90-100 at top
    for i, c in enumerate(counts):
        ax.text(c + 0.05, i, str(c), va="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(path, dpi=130)
    plt.close(fig)


def chart_per_kind_population(per_kind: list[dict[str, Any]], path: Path) -> None:
    """Horizontal bar — how many alive agents per genome kind."""
    if not per_kind:
        return
    rows = per_kind[:12]
    labels = [r["kind"].replace("_", " ") for r in rows]
    counts = [r["n"] for r in rows]
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh(labels, counts, color="#3498db")
    ax.set_xlabel("Alive agents")
    ax.set_title("Population per genome kind (alive)")
    ax.invert_yaxis()
    for i, c in enumerate(counts):
        ax.text(c + 0.2, i, str(c), va="center", fontsize=9)
    plt.tight_layout()
    plt.savefig(path, dpi=130)
    plt.close(fig)


def chart_recent_campaigns(camps: list[dict[str, Any]], path: Path) -> None:
    """Scatter: best_pnl_usd per campaign over the last 24h, colored by kind."""
    if not camps:
        return
    ts = [datetime.fromisoformat(c["created_at"].replace(" ", "T") + "+00:00") for c in camps]
    pnl = [c["best_pnl_usd"] or 0 for c in camps]
    kinds = [c["kind"] for c in camps]
    unique_kinds = sorted(set(kinds))
    cmap = plt.colormaps["tab10"]
    color_map = {k: cmap(i / max(1, len(unique_kinds) - 1)) for i, k in enumerate(unique_kinds)}
    fig, ax = plt.subplots(figsize=(9, 4.5))
    for k in unique_kinds:
        xs = [t for t, kk in zip(ts, kinds) if kk == k]
        ys = [p for p, kk in zip(pnl, kinds) if kk == k]
        ax.scatter(xs, ys, s=60, label=k, color=color_map[k], edgecolors="black", linewidths=0.4)
    ax.axhline(0, color="grey", lw=0.5)
    ax.set_ylabel("Best PnL ($)")
    ax.set_title("Recent campaign best-PnL by genome kind (last 24h)")
    ax.legend(loc="best", fontsize=7, ncol=2)
    fig.autofmt_xdate()
    plt.tight_layout()
    plt.savefig(path, dpi=130)
    plt.close(fig)


def chart_pipeline_diagram(path: Path) -> None:
    """Text-based architecture diagram rendered to PNG."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.axis("off")
    boxes = [
        (0.05, 0.80, 0.20, 0.12, "Polymarket Gamma\n+ CLOB API",          "#3498db"),
        (0.05, 0.55, 0.20, 0.12, "Coinbase WS\n+ REST",                    "#3498db"),
        (0.32, 0.67, 0.18, 0.13, "scanners\n(scan-*.ts)",                 "#1abc9c"),
        (0.32, 0.42, 0.18, 0.13, "market_snapshots\n(SQLite)",            "#9b59b6"),
        (0.57, 0.67, 0.18, 0.13, "Backtest engine\n(scripts/backtest)",   "#1abc9c"),
        (0.57, 0.42, 0.18, 0.13, "Arena evolution\n(factory loops)",      "#e67e22"),
        (0.57, 0.18, 0.18, 0.13, "paper_agents\n(SQLite)",                "#9b59b6"),
        (0.80, 0.55, 0.18, 0.13, "Capsules\n(promotion)",                 "#e74c3c"),
        (0.80, 0.30, 0.18, 0.13, "Live trades\n(Polymarket CLOB)",        "#e74c3c"),
    ]
    for x, y, w, h, text, color in boxes:
        rect = plt.Rectangle((x, y), w, h, facecolor=color, edgecolor="black", alpha=0.85)
        ax.add_patch(rect)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=9, color="white", fontweight="bold")
    arrows = [
        (0.25, 0.86, 0.32, 0.74),  # gamma -> scanners
        (0.25, 0.61, 0.32, 0.55),  # coinbase -> snapshots
        (0.50, 0.73, 0.57, 0.73),  # scanners -> backtest
        (0.50, 0.48, 0.57, 0.48),  # snapshots -> arena
        (0.66, 0.65, 0.66, 0.55),  # backtest -> arena
        (0.66, 0.40, 0.66, 0.31),  # arena -> paper_agents
        (0.75, 0.48, 0.80, 0.60),  # arena -> capsules
        (0.89, 0.55, 0.89, 0.43),  # capsules -> live
    ]
    for x0, y0, x1, y1 in arrows:
        ax.annotate("", xy=(x1, y1), xytext=(x0, y0),
                    arrowprops=dict(arrowstyle="->", color="black", lw=1.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_title("Data flow: ingestion → snapshots → arena → capsules → live", fontsize=11)
    plt.tight_layout()
    plt.savefig(path, dpi=130)
    plt.close(fig)


def chart_timing_lanes(path: Path) -> None:
    """Swimlane-style chart of subsystem cadences."""
    fig, ax = plt.subplots(figsize=(10, 4.5))
    rows = [
        ("worker:research", 60, "#3498db", "scans markets, snapshots prices"),
        ("worker:snapshot",  5, "#3498db", "rapid market snapshot pulls"),
        ("worker:realtime",  1, "#3498db", "WebSocket fills"),
        ("worker:arena", 5, "#1abc9c", "arena tick / generation step"),
        ("factory:btc-5m", 360, "#e67e22", "FAST 6h, DEEP 24h, CHAMP 24h"),
        ("factory:multi", 360, "#e67e22", "per-kind FAST 6h, DEEP 24h"),
        ("worker:graduate", 60, "#9b59b6", "scan + emit graduation-eligible"),
        ("worker:stake-promoter", 240, "#9b59b6", "phase advancement ($2→$5→$10→$20)"),
        ("supervisor", 5, "#7f8c8d", "heartbeat / recovery"),
    ]
    ys = list(range(len(rows)))
    labels = [r[0] for r in rows]
    cadences_min = [r[1] for r in rows]
    colors = [r[2] for r in rows]
    descs = [r[3] for r in rows]
    ax.barh(ys, cadences_min, color=colors, edgecolor="black", alpha=0.85)
    ax.set_yticks(ys, labels)
    ax.invert_yaxis()
    ax.set_xscale("log")
    ax.set_xlabel("Cadence (minutes between runs, log scale)")
    for i, (c, d) in enumerate(zip(cadences_min, descs)):
        ax.text(c * 1.05, i, f"  {d}", va="center", fontsize=8)
    ax.set_title("Subsystem cadences (lower = more frequent)")
    plt.tight_layout()
    plt.savefig(path, dpi=130)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Docx helpers

def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    if level == 0:
        for run in p.runs:
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(0x14, 0x14, 0x14)


def add_paragraph(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = italic
    run.bold = bold


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_i, row in enumerate(rows, start=1):
        row_cells = table.rows[r_i].cells
        for c_i, val in enumerate(row):
            row_cells[c_i].text = str(val) if val is not None else "-"
            for p in row_cells[c_i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacing after table


def add_image(doc: Document, path: Path, *, caption: str = "", width: float = 6.5) -> None:
    if not path.exists():
        add_paragraph(doc, f"[chart missing: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(f"Figure: {caption}")
        run.italic = True
        run.font.size = Pt(9)


def format_bounds(bounds: list[float] | list[str]) -> str:
    if bounds == []:
        return "(populated by seed list)"
    if all(isinstance(x, str) for x in bounds):
        return ", ".join(bounds)
    lo, hi = bounds
    return f"[{lo}, {hi}]"


# ---------------------------------------------------------------------------
# Document 1: Variables + Strategies reference

def build_doc_variables_and_strategies() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "PolymarketAutomation — Agent Variables & Strategies Reference", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")
    run.italic = True
    doc.add_page_break()

    # Overview
    add_heading(doc, "1. Overview", level=1)
    add_paragraph(doc,
        "This document is the single-source reference for every variable an arena agent uses, "
        "every environment knob that controls training and promotion, and the catalogue of all 13 "
        "strategy 'kinds' (genomes) implemented in the codebase. It is mechanically generated from "
        "src/lib/arena/genome.ts (parameter bounds) and src/lib/factory/* (env-driven config). The "
        "factory loops randomly sample parameters inside the listed bounds; evolution then mutates "
        "them around top performers."
    )
    add_paragraph(doc,
        "Read order: §2 explains the fitness formula that ranks agents. §3 lists every env var the "
        "operator can tune. §4 is one section per strategy kind with full parameter tables. §5 is a "
        "cross-reference matrix showing which kinds use which parameters."
    )

    # Fitness formula
    add_heading(doc, "2. Fitness formula and promotion gates", level=1)
    add_paragraph(doc,
        "Two modes are supported, selected via ARENA_FITNESS_MODE."
    )
    add_paragraph(doc, "Mode 'pnl_dd' (legacy):", bold=True)
    add_code(doc, "fitness = pnl_pct − 2.0 × max_dd_pct + activity_bonus")

    add_paragraph(doc, "Mode 'winrate' (default, biased toward 90 %+):", bold=True)
    add_code(doc, "fitness = win_rate^ARENA_WINRATE_POWER × log1p(trades_count) × (1 + pnl_pct) − 0.5 × max_dd_pct")
    add_paragraph(doc,
        "Agents with fewer than ARENA_MIN_TRADES_FOR_RANKING (default 30) closed trades return a "
        "sentinel score so a 3-trade lucky streak can't win a generation. The squared win-rate term "
        "(default exponent 2) steepens the gradient near 1.0 so a 90 % winner outscores a 60 % winner "
        "by 2.25× rather than 1.5×."
    )

    add_paragraph(doc, "Promotion ladder:", bold=True)
    add_paragraph(doc, "sim → paper → live_eligible → live")
    add_paragraph(doc, "Gates at each boundary:")
    add_table(doc,
        ["Boundary", "Min trades", "Min win rate", "Min PnL", "Other"],
        [
            ["sim → paper",
             "GRADUATION_MIN_TRADES = 15",
             "GRADUATION_MIN_WIN_RATE = 0.90",
             "GRADUATION_MIN_PNL_USD = $10",
             "auto-graduates campaign + consistent-winner agents"],
            ["paper → live",
             "ARENA_AUTO_PROMOTE_MIN_TRADES = 3",
             "ARENA_AUTO_PROMOTE_MIN_WIN_RATE = 0.90",
             "MIN_LIVE_CAPSULE_PNL_USD = $96",
             "ALLOW_AUTO_PROMOTE=1 required, correlation veto, kind safety ceiling"],
        ]
    )

    # Env vars
    add_heading(doc, "3. Environment variable reference", level=1)
    add_paragraph(doc,
        "Every knob below can be set in .env.local or exported in the shell that launches the workers. "
        "Defaults are what the running factories use unless overridden."
    )
    add_table(doc,
        ["Variable", "Default", "Purpose"],
        [[v["name"], v["default"], v["purpose"]] for v in ENV_VARS],
    )

    # Strategy catalogue
    add_heading(doc, "4. Strategy catalogue (13 genome kinds)", level=1)
    add_paragraph(doc,
        "Each subsection covers one kind. Param tables show the parameter name, its range (or "
        "enumerated values for discrete params), and what the parameter controls. The factories "
        "sample uniformly inside each range when seeding random variants; evolution mutates around "
        "the best-performing agents within a per-pct deviation."
    )

    for s in STRATEGIES:
        add_heading(doc, f"4.{STRATEGIES.index(s) + 1}. {s['title']} ({s['kind']})", level=2)
        add_paragraph(doc, f"Venue: {s['venue']}", italic=True)
        add_paragraph(doc, f"Live-eligible by default: {s['live_eligible']}", italic=True)
        if s.get("note"):
            add_paragraph(doc, f"Note: {s['note']}", italic=True)
        add_paragraph(doc, s["summary"])
        add_paragraph(doc, f"Fires when: {s['fires_when']}", bold=True)
        rows = []
        for name, (bounds, desc) in s["params"].items():
            rows.append([name, format_bounds(bounds), desc])
        add_table(doc, ["Parameter", "Range / values", "Controls"], rows)

    # Cross-reference matrix
    add_heading(doc, "5. Parameter cross-reference matrix", level=1)
    add_paragraph(doc,
        "Rows are common parameter names; columns are strategy kinds. ✓ means the kind uses that "
        "parameter. Useful for the operator who wants to know 'which strategies care about lookback_h?'"
    )
    common_params = sorted({
        p for s in STRATEGIES for p in s["params"].keys()
    })
    headers = ["Parameter"] + [s["kind"] for s in STRATEGIES]
    matrix_rows = []
    for p in common_params:
        row = [p] + ["✓" if p in s["params"] else "" for s in STRATEGIES]
        matrix_rows.append(row)
    add_table(doc, headers, matrix_rows)

    doc.save(OUT_VARS)
    print(f"  -> {OUT_VARS} ({OUT_VARS.stat().st_size // 1024} KB)")


# ---------------------------------------------------------------------------
# Document 2: Architecture + Timing

def build_doc_architecture_and_timing(stats: dict[str, Any], state: dict[str, Any]) -> None:
    # Render all charts first so the docx writer can embed them.
    if not stats.get("db_missing"):
        chart_win_rate_distribution(stats["win_rate_distribution"], CHARTS / "win-rate-distribution.png")
        chart_per_kind_population(stats["per_kind"], CHARTS / "per-kind-population.png")
        chart_recent_campaigns(stats["recent_campaigns"], CHARTS / "recent-campaigns.png")
    chart_pipeline_diagram(CHARTS / "pipeline.png")
    chart_timing_lanes(CHARTS / "timing-lanes.png")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "PolymarketAutomation — Architecture & Strategic Timing Map", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  ·  live snapshot")
    run.italic = True
    doc.add_page_break()

    # Exec summary
    add_heading(doc, "Executive summary", level=1)
    if stats.get("db_missing"):
        add_paragraph(doc, "DB not available — operational stats sections skipped.")
    else:
        best_pct = stats["best_win_rate"] * 100
        add_paragraph(doc,
            f"At time of generation the population holds {stats['alive_count']} alive agents and "
            f"{stats['dead_count']} retired. {stats['qualifying_count']} agents have cleared the "
            f"30-trade ranking floor. Best win rate among qualifying agents: {best_pct:.1f} %. The "
            f"system holds {stats['snapshot_total']:,} market snapshots collected between "
            f"{stats['snapshot_earliest']} and {stats['snapshot_latest']}."
        )
    if not state.get("missing"):
        for name, f in state.get("factories", {}).items():
            add_paragraph(doc,
                f"Factory '{name}': desired={f.get('desired')}, pid={f.get('pid')}, "
                f"started_at={f.get('startedAt') or '-'}, restarts={f.get('startCount')}."
            )

    # Architecture
    add_heading(doc, "1. System architecture", level=1)
    add_paragraph(doc,
        "Six layers, each owned by a distinct module and worker:"
    )
    add_table(doc,
        ["Layer", "Owns", "Workers / scripts", "Persistence"],
        [
            ["Ingestion", "Polymarket Gamma + CLOB, Coinbase REST + WS",
             "worker:research, worker:realtime, worker:snapshot, scan:*",
             "market_snapshots, paper_trades"],
            ["Snapshot store", "OHLC + orderbook + price-history per token",
             "snapshot:cb-depth, snapshot:cb-trades, snapshot:cb-stats",
             "polymarket.db, historical-candles.db"],
            ["Backtesting", "Replay snapshots through decision functions",
             "scripts/backtest, src/lib/backtest/engine",
             "performance_metrics, training_campaigns"],
            ["Arena", "Evolutionary search over genome kinds + params",
             "factory:btc-5m, factory:multi, worker:arena, worker:graduate",
             "paper_agents, training_campaigns, candidates"],
            ["Capsules", "Promotion to paper → live with risk caps",
             "worker:stake-promoter, arena:promote-live",
             "capsules"],
            ["Live execution", "Polymarket CLOB v2 fills, reconcile",
             "worker:reconcile, worker:realtime, packages/adapters/polymarket",
             "paper_trades (with venue=live)"],
        ]
    )
    add_image(doc, CHARTS / "pipeline.png",
              caption="Data flow from venue feeds through snapshots, backtest, arena, capsules to live.")

    # Timing
    add_heading(doc, "2. Subsystem cadences", level=1)
    add_paragraph(doc,
        "The cadences below are documented defaults — override via env or --interval-* flags on "
        "each worker. Multiple workers can run in parallel; the supervisor ensures heartbeats."
    )
    add_table(doc,
        ["Subsystem", "Default cadence", "What runs"],
        [
            ["supervisor",            "every 5 min (cron)",   "stale-heartbeat detection, recovery"],
            ["worker:realtime",       "WebSocket (~1s)",       "live fills, real-time book updates"],
            ["worker:snapshot",       "5 min",                 "rapid snapshot pulls during research-loop windows"],
            ["worker:research",       "~hourly",               "market sweep, signal evaluation, version proposals"],
            ["worker:arena",          "5 min tick (10 min backoff at trade cap)", "arena step + generation evolve"],
            ["factory:btc-5m",        "fast 6 h / deep 24 h / champ 24 h",        "consistent-winner BTC 5m binary directional"],
            ["factory:multi",         "fast 6 h / deep 24 h per kind (12 kinds)", "all non-BTC-5m genome kinds"],
            ["worker:graduate",       "60 min",                "emit graduation-eligible for paper capsules"],
            ["worker:stake-promoter", "4 h",                   "advance phases $2 → $5 → $10 → $20"],
            ["worker:reconcile",      "~5 min",                "venue + capsule state diff"],
        ]
    )
    add_image(doc, CHARTS / "timing-lanes.png",
              caption="Cadence per subsystem on log scale — lower = more frequent.")

    # Live stats
    if not stats.get("db_missing"):
        add_heading(doc, "3. Current population (live)", level=1)
        add_paragraph(doc,
            f"{stats['alive_count']} alive agents, {stats['qualifying_count']} of which have "
            f"cleared the 30-trade ranking floor."
        )
        add_image(doc, CHARTS / "win-rate-distribution.png",
                  caption=f"Win-rate distribution — current best {stats['best_win_rate'] * 100:.1f}% vs 90% gate.")

        add_paragraph(doc, "Population by genome kind:", bold=True)
        add_table(doc,
            ["Kind", "Alive count", "Avg trades", "Avg PnL"],
            [
                [r["kind"], r["n"], f"{(r['avg_t'] or 0):.1f}", f"${(r['avg_pnl'] or 0):.2f}"]
                for r in stats["per_kind"]
            ],
        )
        add_image(doc, CHARTS / "per-kind-population.png",
                  caption="Alive population by genome kind — uneven coverage = work for the multi-factory.")

        add_paragraph(doc, "Top 10 alive agents (≥30 trades, sorted by win rate):", bold=True)
        add_table(doc,
            ["#", "Name", "Kind", "Trades", "Wins", "Win rate", "PnL"],
            [
                [i + 1, r["name"][:38], r["kind"][:24], r["trades_count"], r["wins_count"],
                 f"{(r['wins_count']/r['trades_count'])*100:.1f}%", f"${r['realized_pnl_usd']:.2f}"]
                for i, r in enumerate(stats["top_agents"])
            ],
        )

        add_heading(doc, "4. Recent factory output (last 24h)", level=1)
        add_paragraph(doc,
            f"{len(stats['recent_campaigns'])} training campaigns logged in the last 24 hours."
        )
        if stats["recent_campaigns"]:
            add_image(doc, CHARTS / "recent-campaigns.png",
                      caption="Best PnL per campaign, by genome kind — positive cluster = the strategies that are working.")
            add_table(doc,
                ["When", "Name", "Kind", "Produced", "Best PnL"],
                [
                    [c["created_at"], c["name"][:42], c["kind"][:22],
                     c["candidates_produced"] or 0,
                     f"${c['best_pnl_usd']:.2f}" if c["best_pnl_usd"] is not None else "-"]
                    for c in stats["recent_campaigns"][:15]
                ],
            )

    # Strategic positioning
    add_heading(doc, "5. Strategic positioning — where each strategy fits", level=1)
    add_paragraph(doc,
        "Each strategy occupies a distinct slot in the timing + signal space. The matrix below shows "
        "the time horizon a strategy operates over, the venue, and the structural advantage it claims."
    )
    add_table(doc,
        ["Strategy", "Time horizon", "Venue", "Edge thesis", "Capacity"],
        [
            ["poly_short_binary_directional", "5–15 min",  "Polymarket", "Polymarket reprices slower than Coinbase by ~2.7s", "High — many windows/day"],
            ["polymarket_market_maker",       "continuous","Polymarket", "Spread capture as inventory rotates",                  "Moderate — limited by depth"],
            ["llm_probability_oracle",        "minutes",   "Polymarket", "LLM beats midpoint when news is in-context",            "Low — LLM cost-bound"],
            ["poly_fade_spike",               "hours",     "Polymarket", "Mean reversion after attention bursts",                "Moderate"],
            ["poly_breakout",                 "hours",     "Polymarket", "Continuation after attention bursts",                  "Moderate"],
            ["cross_venue_arb",               "hours",     "Poly + CB", "BS-implied vs Polymarket-implied mispricing",          "Low — narrow opportunity window"],
            ["cb_breakout",                   "30 min – 24 h", "Coinbase", "Range break-out on liquid crypto",                    "High"],
            ["cb_mean_reversion",             "hours",     "Coinbase",  "Reversion from z-score extremes",                       "High"],
            ["cb_momentum_burst",             "5–120 min", "Coinbase",  "Acceleration spikes (news, halt-reopen)",              "Moderate"],
            ["category_specialist",           "varies",    "Polymarket","Outperformance restricted to one tag",                 "Low — one slice at a time"],
            ["wallet_copy_filtered",          "minutes",   "Polymarket","Mirror a proven wallet's edge",                        "Depends on source"],
            ["multi_strategy",                "composite", "Either",    "Stacked sub-strategies with priority selection",        "Moderate"],
            ["random_walk_baseline",          "any",       "Polymarket","Control — any real edge must beat this",                "N/A"],
        ]
    )

    add_heading(doc, "6. Gaps and recommended next moves", level=1)
    add_paragraph(doc,
        "Things the architecture supports but isn't actively using:"
    )
    add_paragraph(doc,
        "• 5-min BTC Up/Down market discovery — the snapshot worker is currently pulling event-resolution "
        "markets, not the high-frequency crypto binary contracts that the BTC-5m factory was tuned for. "
        "Adding a discovery query against the Gamma API filtering for crypto 5/15-minute binaries would "
        "give the BTC-5m factory the data it needs to converge.")
    add_paragraph(doc,
        "• Markov-persistence strategy seeding — the strategy code is in src/lib/strategies/ but no genome "
        "kind has been registered for it. Once registered, multi-factory will breed it like any other kind.")
    add_paragraph(doc,
        "• Maker-only execution path — wired into packages/core/src/venue/router.ts but not yet enabled by "
        "default. Becker's data shows +1.12 % maker / −1.12 % taker, so flipping this on is roughly worth "
        "+2.24 pp on every fill.")
    add_paragraph(doc,
        "• Live multi-coin fan-out — the factories operate on BTC + ETH + SOL via cb_* genomes. Adding XRP, "
        "DOGE, SHIB, etc. is a one-line change in PARAM_BOUNDS but should be paired with snapshot coverage.")
    add_paragraph(doc,
        "• Win-rate progression telemetry — the dashboard tracks the 'best win rate' over the dashboard's "
        "uptime. Persisting that into a daily snapshots table would give a 30-day chart instead of a 24h "
        "linear projection.")

    add_heading(doc, "7. What to watch", level=1)
    add_paragraph(doc,
        "Three signals indicate whether the system is converging toward the 90 % win-rate goal."
    )
    add_table(doc,
        ["Signal", "Source", "Healthy", "Stuck"],
        [
            ["Best win rate (Δ per day)", "factory dashboard",  "+1 to +3 pp/day", "flat or negative for 3+ days"],
            ["Population in 70-80 % bucket", "factory dashboard", "growing weekly", "stays at 0 for 7+ days"],
            ["Factory campaign best PnL", "training_campaigns", "consistently > $0", "consistently < $0 across kinds"],
        ]
    )

    doc.save(OUT_ARCH)
    print(f"  -> {OUT_ARCH} ({OUT_ARCH.stat().st_size // 1024} KB)")


# ---------------------------------------------------------------------------
# Main

def build_doc_latency_and_event_map() -> None:
    """Doc 3: synthesizes the polymarket-2dollar-bot/mac framework + HFT
    docs/strategies/ into a PolymarketAutomation-specific latency + event
    timing playbook. Cross-references the source files in both repos.
    """
    chart_pipeline_diagram(CHARTS / "pipeline.png")
    chart_timing_lanes(CHARTS / "timing-lanes.png")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "PolymarketAutomation — Latency + Event Map", level=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}  ·  "
        "Synthesized from polymarket-2dollar-bot/mac + HFT/docs/strategies"
    )
    run.italic = True
    doc.add_page_break()

    add_heading(doc, "Executive summary", level=1)
    add_paragraph(doc,
        "This document is the operator's reference for thinking about latency, "
        "event lifecycle, and signal freshness in the PolymarketAutomation arena. "
        "It synthesizes two existing source frameworks — the polymarket-2dollar-bot "
        "'mac' branch (live WS pricefeed + microstructure scanners + architecture "
        "report) and the HFT repo's strategy docs (latency-arbitrage, event-driven, "
        "microstructure-signals) — into a playbook applicable to the arena agents "
        "we breed in this repo."
    )
    add_paragraph(doc,
        "Three new genome kinds and one TickContext-adjacent timing module land "
        "as part of this work: markov_persistence (timing-aware), "
        "poly_arbitrage_set (risk-free YES+NO buy), and poly_repricing "
        "(directional spot-vs-market gap)."
    )

    # ── Latency tiers
    add_heading(doc, "1. Latency tiers (T0–T4)", level=1)
    add_paragraph(doc,
        "From HFT/docs/strategies/latency-arbitrage.md §2.1 — what 'time' "
        "means in a leader-lagger race. Every variant has the same shape: "
        "read the leader → predict where the lagger will be → trade against "
        "the lagger's stale quote → close when the lagger updates."
    )
    add_table(doc,
        ["Tier", "Budget", "Reachable from", "Examples"],
        [
            ["T0", "µs (wire)",        "microwave + colo, FPGA",     "ES → SPY (Eurex 3.2ns = €75M/yr)"],
            ["T1", "µs (software)",    "kernel-bypass NIC, single-thread C++", "CME futures → cash, BiNance lagger"],
            ["T2", "10–100 ms",        "same-region cloud, no colo", "Binance USDM → Coinbase spot"],
            ["T3", "100–1000 ms",      "different region, REST polling",     "** PolymarketAutomation's tier **"],
            ["T4", "1–15 s (block)",   "blockchain block latency",  "DEX-CEX stale-oracle arb"],
        ]
    )
    add_paragraph(doc,
        "PolymarketAutomation operates in T3. Our snapshot worker polls; our "
        "realtime worker uses WS (sub-second when healthy) for crypto ticks "
        "but still goes through HTTPS for order placement. We cannot win a "
        "race against firms in T0–T1, but the leader-lagger gap on Polymarket "
        "is 2.7–12 seconds — well above our floor.",
    )

    # ── Binary lifecycle
    add_heading(doc, "2. Polymarket binary lifecycle", level=1)
    add_paragraph(doc,
        "Reference: HFT event-driven.md §2.6 + polymarket-2dollar-bot "
        "analysis/architecture_report.py. Each 5-min and 15-min binary has a "
        "deterministic lifecycle from window-open to settlement:"
    )
    add_table(doc,
        ["Phase", "Range", "Strategic signal"],
        [
            ["pre-window",  "now < expiry − duration",        "market exists but window hasn't started (rare 5-min)"],
            ["opening",     "0 % – 25 % of window elapsed",   "signal forming, Markov persistence is unstable"],
            ["mid-window",  "25 % – 75 %",                    "signal established, edge widest, Markov mature"],
            ["late-window", "75 % – 100 % minus cutoff",      "sample-size dominated, exit-liq premium grows"],
            ["post-cutoff", "minToResolution ≤ cutoff (~3m)", "Polymarket order-book lockout — no new orders"],
            ["resolved",    "expiry passed",                  "settled at oracle close"],
        ]
    )
    add_paragraph(doc,
        "Implementation: src/lib/arena/event-timing.ts — `eventPhase({expiryIso, "
        "durationMin, now, cutoffMin})`. Pure function. Phase split at 0.25/0.75 "
        "matches the polymarket-2dollar-bot architecture report."
    )

    # ── Microstructure signals
    add_heading(doc, "3. Microstructure signals", level=1)
    add_paragraph(doc,
        "Ported from polymarket-2dollar-bot/polybot/microstructure.py and "
        "HFT/docs/strategies/microstructure-signals.md §2.3 (OFI):"
    )
    add_table(doc,
        ["Signal", "Returns", "PolymarketAutomation home", "Use"],
        [
            ["arbitrageEdge",
             "Opportunity | null",
             "src/lib/quant/microstructure.ts",
             "BOTH-side buy when YES+NO+fees < $1; locked profit"],
            ["directionalArbTilt",
             "Opportunity | null",
             "src/lib/quant/microstructure.ts",
             "Arb base + model view → tilt to under-priced side"],
            ["nearResolutionEdge",
             "Opportunity | null",
             "src/lib/quant/microstructure.ts",
             "Buy late-window almost-certain side at 0.95–0.99"],
            ["orderbookImbalance",
             "number ∈ [-1, 1]",
             "src/lib/quant/microstructure.ts",
             "Top-N depth skew — bid-heavy / ask-heavy lag signal"],
            ["repricingEdge",
             "Opportunity | null",
             "src/lib/quant/microstructure.ts",
             "Fair P(YES) vs market gap — directional bet on lag"],
        ]
    )

    # ── Genome enhancements
    add_heading(doc, "4. Agent enhancements landed", level=1)
    add_paragraph(doc,
        "Three genome kinds now use the latency/event framework:"
    )
    add_table(doc,
        ["Genome kind", "Strategy", "Timing-aware?", "Source port"],
        [
            ["markov_persistence",
             "p(j*,j*) ≥ 0.87 + Markov MC + Becker calibration",
             "Yes — 4 timing-gate params (min/max time-to-resolution, phase filter, signal freshness)",
             "@0xRicker / @de1lymoon articles"],
            ["poly_arbitrage_set",
             "Buy YES + NO when ask sum + fees < $1",
             "No — fires on book inefficiency, not timing",
             "polymarket-2dollar-bot arbitrage_edge()"],
            ["poly_repricing",
             "Spot-vs-market gap via Coinbase tick → fair P(YES)",
             "Yes — same 4 timing-gate params as markov",
             "polymarket-2dollar-bot repricing_edge()"],
        ]
    )

    # ── New genome params
    add_heading(doc, "5. New genome params (timing gates)", level=1)
    add_paragraph(doc,
        "These appear on markov_persistence and poly_repricing. randomGenome "
        "samples them by default; each is configurable per agent via the "
        "ARENA tuning interfaces."
    )
    add_table(doc,
        ["Param", "Range", "Default behavior (disabled)", "What it gates"],
        [
            ["min_time_to_resolution_min", "[0, 30]", "0", "Skip if minToResolution < this (avoid cutoff zone)"],
            ["max_time_to_resolution_min", "[1, 999]", "999", "Skip if minToResolution > this (skip event markets)"],
            ["event_phase_filter",         "enum",    "any", "Restrict to specific lifecycle phase"],
            ["max_signal_age_sec",         "[1, 9999]", "9999", "Skip if Coinbase WS tick older than this"],
        ]
    )

    # ── Pipeline overview
    add_heading(doc, "6. Data pipeline (where freshness lives)", level=1)
    add_image(doc, CHARTS / "pipeline.png",
              caption="Coinbase + Polymarket feeds → snapshot/realtime stores → arena → capsules → live.")
    add_paragraph(doc,
        "Critical freshness checkpoints, in order:"
    )
    add_paragraph(doc,
        "• worker:realtime — sub-second Polymarket + Coinbase WS, 1-sec debounce, writes realtime_ticks."
        " Equivalent to polymarket-2dollar-bot's pricefeed.py + scripts/pricefeed.py.")
    add_paragraph(doc,
        "• worker:snapshot — periodic deeper poll (orderbook, history). Equivalent to scripts/scan*.py in 2dollar-bot.")
    add_paragraph(doc,
        "• decideMarkovPersistence + decidePolyRepricing — read latestRealtimeTicks() at decision time. "
        "If the most recent Coinbase tick is older than max_signal_age_sec, the strategy holds. "
        "This is the same pattern as 2dollar-bot's `latest()` check in pricefeed.py.")
    add_paragraph(doc,
        "• Execution: when the agent fires, the executor reads a SECOND fresh tick to ask-re-check. "
        "Same pattern as 2dollar-bot's commit 581980c \"freshest data at decision + execution\".")

    add_image(doc, CHARTS / "timing-lanes.png",
              caption="Per-subsystem cadence (log scale) — supervisor every 5 min, stake-promoter every 4 hours.")

    # ── What to watch
    add_heading(doc, "7. Strategic positioning per kind", level=1)
    add_table(doc,
        ["Kind", "Best phase", "Tick-freshness sensitivity", "Capacity"],
        [
            ["markov_persistence",        "mid-window",    "high (5-15s)",  "moderate"],
            ["poly_arbitrage_set",        "any",           "low",            "low — rare"],
            ["poly_repricing",            "mid-or-late",   "very high (<5s)", "moderate"],
            ["poly_short_binary_directional", "post-cutoff peek", "high",  "high — main BTC-5m"],
            ["polymarket_market_maker",   "continuous",    "low",            "low — depth-bound"],
        ]
    )

    # ── Source cross-reference
    add_heading(doc, "8. Source cross-reference", level=1)
    add_paragraph(doc, "Where each concept lives in each repo:", bold=True)
    add_table(doc,
        ["Concept", "polymarket-2dollar-bot/mac", "HFT/docs/strategies", "PolymarketAutomation"],
        [
            ["Sub-second WS price store",
             "polybot/pricefeed.py + scripts/pricefeed.py",
             "—",
             "src/lib/arena/realtime-ticks.ts"],
            ["Arbitrage edge",
             "polybot/microstructure.py arbitrage_edge()",
             "—",
             "src/lib/quant/microstructure.ts"],
            ["Repricing edge",
             "polybot/microstructure.py repricing_edge()",
             "—",
             "src/lib/quant/microstructure.ts + decidePolyRepricing"],
            ["OFI / OBI",
             "polybot/microstructure.py orderbook_imbalance()",
             "microstructure-signals.md §2.3 (OFI formal)",
             "src/lib/quant/microstructure.ts orderbookImbalance()"],
            ["Latency tier framework",
             "—",
             "latency-arbitrage.md §2.1",
             "src/lib/arena/event-timing.ts (T3 implicit)"],
            ["Binary lifecycle phases",
             "analysis/architecture_report.py",
             "event-driven.md §2.6",
             "src/lib/arena/event-timing.ts eventPhase()"],
            ["Markov persistence threshold",
             "—",
             "—",
             "src/lib/quant/markov.ts persistenceProbability()"],
            ["Becker calibration",
             "polybot/quant.py becker_*",
             "—",
             "src/lib/quant/becker-calibration.ts"],
        ]
    )

    add_heading(doc, "9. Next moves", level=1)
    add_paragraph(doc,
        "Items in the original polymarket-2dollar-bot framework not yet ported:"
    )
    add_paragraph(doc,
        "• directional_arb_tilt as its own genome kind (helper exists in microstructure.ts; needs decide function).")
    add_paragraph(doc,
        "• near_resolution_edge as its own genome kind (helper exists; near-resolution-scrape.ts has the strategy already, would need genome wrapper).")
    add_paragraph(doc,
        "• Hermes-style nightly self-tuner (2dollar-bot's tuner.py) — currently we have arena evolution but no per-agent param-update loop driven by realized PnL diagnostics.")
    add_paragraph(doc,
        "• Walk-forward / PBO overfit battery (HFT's commit cd85cad/289dafb pattern) — important before any live promotion at scale.")

    doc.save(OUT_LATENCY)
    print(f"  -> {OUT_LATENCY} ({OUT_LATENCY.stat().st_size // 1024} KB)")


def main() -> None:
    print(f"Generating docs into {ANALYSIS}")
    print("Document 1: Variables + Strategies reference")
    build_doc_variables_and_strategies()
    print("Document 2: Architecture + Timing report")
    stats = fetch_live_stats()
    state = read_factory_state()
    build_doc_architecture_and_timing(stats, state)
    print("Document 3: Latency + Event Map")
    build_doc_latency_and_event_map()
    print("Done.")


if __name__ == "__main__":
    main()
